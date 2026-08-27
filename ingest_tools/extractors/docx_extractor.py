"""DOCX extractor using python-docx (tables + paragraphs)."""

import logging
from typing import Optional

from extractors.base_extractor import BaseExtractor, ExtractionResult

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """Extract text and tables from Word (.docx) documents.

    Strategy:
        1. Extract structured tables (each table row -> list of cell strings).
        2. Also collect plain paragraph text as fallback.
    """

    def extract(
        self,
        file_path: str,
        pages: Optional[list[int]] = None,
    ) -> ExtractionResult:
        self._validate_file(file_path)
        logger.info(f"Extracting DOCX: {file_path}")

        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required to extract .docx files")

        result = ExtractionResult()
        doc = docx.Document(file_path)
        result.pages = 1
        result.metadata["format"] = "docx"

        for table in doc.tables:
            cleaned = self._table_to_list(table)
            if cleaned:
                result.tables.append(cleaned)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        result.raw_text = "\n".join(text_parts)

        logger.info(
            f"Extracted {len(result.tables)} table(s), "
            f"{len(result.raw_text)} chars of text"
        )
        return result

    def _table_to_list(self, table) -> list[list[str]]:
        cleaned = []
        seen_rows = set()
        for row in table.rows:
            cells = []
            for cell in row.cells:
                val = cell.text.replace("\n", " ").strip()
                cells.append(val)
            key = tuple(cells)
            if key in seen_rows:
                continue
            seen_rows.add(key)

            cleaned_row = [c for c in cells]
            if any(cleaned_row):
                cleaned.append(cleaned_row)

        return cleaned if len(cleaned) >= 2 else []
